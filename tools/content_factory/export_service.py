import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from fpdf import FPDF
from content_factory_service import coursewareDocument

class ExportService:
    def __init__(self, output_dir: Path):
        self.output_dir = output_dir
        self.output_dir.mkdir(parents=True, exist_ok=True)
        # Theme Colors
        self.color_cyan = (0, 255, 255)
        self.color_purple = (147, 0, 211)
        self.color_bg = (10, 10, 40) # Dark Navy

    def to_docx(self, doc: coursewareDocument, filename: str):
        """Exports the courseware to a Microsoft Word document."""
        path = self.output_dir / filename
        document = Document()
        
        # Title
        document.add_heading(doc.title, 0)
        document.add_paragraph(doc.summary, style='Intense Quote')

        for section in doc.sections:
            document.add_heading(section.title, level=1)
            
            # Integrator Track
            document.add_heading('[INTEGRATOR] Track', level=2)
            document.add_paragraph(section.integrator_content)
            
            # Architect Track
            document.add_heading('[ARCHITECT] Track', level=2)
            document.add_paragraph(section.architect_content)
            
            # Diagrams
            if section.diagrams:
                document.add_heading('Visual Assets', level=2)
                for diag in section.diagrams:
                    document.add_paragraph(diag.title, style='List Bullet')
                    document.add_paragraph(f"Prompt: {diag.nano_banana_prompt}")
                    # If the image exists on disk, we can embed it
                    # Note: In process_courseware.py we know the path
                    # This service just handles the structure.
            
            document.add_page_break()

        document.save(path)
        print(f"[OK] DOCX saved to {path}")

    def to_pdf(self, doc: coursewareDocument, filename: str):
        """Exports the courseware to a PDF document."""
        path = self.output_dir / filename
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        
        # Add Page
        pdf.add_page()
        
        # Title
        pdf.set_font("Helvetica", "B", 24)
        pdf.set_text_color(*self.color_cyan)
        pdf.cell(0, 20, doc.title, ln=True, align='C')
        
        # Summary
        pdf.set_font("Helvetica", "I", 12)
        pdf.set_text_color(100, 100, 100)
        pdf.multi_cell(0, 10, doc.summary)
        pdf.ln(10)

        for section in doc.sections:
            # Section Title
            pdf.set_font("Helvetica", "B", 18)
            pdf.set_text_color(*self.color_purple)
            pdf.cell(0, 15, section.title, ln=True)
            
            # Integrator
            pdf.set_font("Helvetica", "B", 14)
            pdf.set_text_color(0, 0, 0)
            pdf.cell(0, 10, "[INTEGRATOR] Track", ln=True)
            pdf.set_font("Helvetica", "", 11)
            pdf.multi_cell(0, 7, section.integrator_content)
            pdf.ln(5)
            
            # Architect
            pdf.set_font("Helvetica", "B", 14)
            pdf.cell(0, 10, "[ARCHITECT] Track", ln=True)
            pdf.set_font("Helvetica", "", 11)
            pdf.multi_cell(0, 7, section.architect_content)
            pdf.ln(10)
            
            # Start new page for next section if it's too long
            # Simple page management
            if pdf.get_y() > 250:
                pdf.add_page()

        pdf.output(str(path))
        print(f"[OK] PDF saved to {path}")
